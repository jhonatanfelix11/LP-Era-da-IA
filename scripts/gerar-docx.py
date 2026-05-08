from docx import Document
from docx.shared import Pt, Mm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# Page size A4
section = doc.sections[0]
section.page_width = Mm(210)
section.page_height = Mm(297)
section.left_margin = Mm(25)
section.right_margin = Mm(25)
section.top_margin = Mm(22)
section.bottom_margin = Mm(22)

# ── helpers ──────────────────────────────────────────────────────────────────

def set_font(run, name="Calibri", size=11, bold=False, italic=False,
             color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1, color=(5, 14, 26)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    sizes = {1: 24, 2: 18, 3: 14, 4: 12}
    set_font(run, "Calibri", sizes.get(level, 12), bold=True, color=color)
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_body(doc, text, italic=False, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, "Calibri", 11, italic=italic)
    if indent:
        p.paragraph_format.left_indent = Mm(8)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(16)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_font(run, "Calibri", 11)
    p.paragraph_format.left_indent = Mm(8 + level * 6)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_callout(doc, text, label="DESTAQUE"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Mm(8)
    p.paragraph_format.right_indent = Mm(8)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    run_label = p.add_run(f"[{label}] ")
    set_font(run_label, "Calibri", 10, bold=True, color=(37, 99, 235))
    run_text = p.add_run(text)
    set_font(run_text, "Calibri", 11, italic=True, color=(30, 58, 138))
    return p

def add_divider(doc):
    p = doc.add_paragraph("─" * 72)
    run = p.runs[0]
    set_font(run, "Calibri", 9, color=(147, 197, 253))
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_stat(doc, number, label):
    p = doc.add_paragraph()
    run_num = p.add_run(number + "  ")
    set_font(run_num, "Calibri", 20, bold=True, color=(37, 99, 235))
    run_lbl = p.add_run(label)
    set_font(run_lbl, "Calibri", 11, color=(75, 85, 99))
    p.paragraph_format.space_after = Pt(6)
    return p

def page_break(doc):
    doc.add_page_break()

# ── CAPA ─────────────────────────────────────────────────────────────────────

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ERA DA I.A.")
set_font(run, "Calibri", 36, bold=True, color=(5, 14, 26))
p.paragraph_format.space_before = Pt(40)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Como Lucrar Agora\nAntes que Seja Tarde Demais")
set_font(run, "Calibri", 22, italic=True, color=(37, 99, 235))

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("O Mapa Completo de Monetização com Inteligência Artificial\npara Profissionais Criativos Brasileiros")
set_font(run, "Calibri", 13, color=(75, 85, 99))

add_divider(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("2025 — Todos os direitos reservados")
set_font(run, "Calibri", 9, color=(107, 114, 128))

page_break(doc)

# ── AVISO LEGAL ──────────────────────────────────────────────────────────────

add_heading(doc, "Aviso Legal", 2)
add_body(doc, "Este material tem caráter exclusivamente educacional e informativo. Os resultados financeiros mencionados são estimativas baseadas em médias de mercado e não constituem garantia de ganhos. Resultados individuais variam conforme dedicação, habilidade e condições de mercado.")
add_body(doc, "© 2025 — Todos os direitos reservados. Proibida a reprodução total ou parcial sem autorização expressa do autor.")

page_break(doc)

# ── SUMÁRIO ───────────────────────────────────────────────────────────────────

add_heading(doc, "Sumário", 1)

chapters = [
    ("Introdução", "O Momento que Muda Carreiras"),
    ("Capítulo 1", "O Mapa do Mercado de I.A. em 2025"),
    ("Capítulo 2", "Os Hubs: Magnific.ai e Higgsfield.ai"),
    ("Capítulo 3", "Os Agentes: Claude, ChatGPT e Gemini"),
    ("Capítulo 4", "Geradores de Imagem: DALL-E 3 e Imagen 3"),
    ("Capítulo 5", "Carreiras e Quanto Cobrar"),
    ("Capítulo 6", "Fluxos de Trabalho Profissionais"),
    ("Capítulo 7", "Seu Plano de 30 Dias"),
    ("Conclusão", "A Janela Está Aberta — Por Enquanto"),
    ("Bônus 1", "20 Prompts Prontos para Usar Hoje"),
    ("Bônus 2", "Tabela de Preços por Serviço"),
    ("Bônus 3", "30 Ferramentas Essenciais de I.A."),
]

for prefix, title in chapters:
    p = doc.add_paragraph()
    run_prefix = p.add_run(f"{prefix}  ")
    set_font(run_prefix, "Calibri", 11, bold=True, color=(37, 99, 235))
    run_title = p.add_run(title)
    set_font(run_title, "Calibri", 11)
    p.paragraph_format.space_after = Pt(4)

page_break(doc)

# ── SOBRE O AUTOR ─────────────────────────────────────────────────────────────

add_heading(doc, "Sobre o Autor", 1)
add_body(doc, "Profissional com mais de 8 anos de experiência em criação digital, tendo trabalhado com marcas nacionais e internacionais nas áreas de design, vídeo e marketing de conteúdo. Acompanhou de perto a revolução das ferramentas de I.A. desde 2022 e passou os últimos 2 anos testando, validando e monetizando cada nova categoria de ferramenta.")
add_body(doc, "Este eBook é o resultado prático dessa jornada: não teoria sobre o futuro da I.A., mas um mapa do que funciona agora, para profissionais brasileiros, com resultados reais.")

page_break(doc)

# ── CARTA AO LEITOR ────────────────────────────────────────────────────────────

add_heading(doc, "Carta ao Leitor", 1)
add_body(doc, "Se você está lendo isso, provavelmente já sentiu aquela sensação desconfortável de ver alguém usando I.A. para fazer em 10 minutos o que você leva horas para fazer — e faturando com isso.")
add_body(doc, "Eu sei exatamente como isso é. Passei por isso em 2022, quando vi as primeiras imagens do Midjourney e pensei: \"isso vai mudar tudo\". E mudou. A diferença é que alguns profissionais entraram cedo, entenderam o mapa, e hoje têm uma vantagem que é difícil de recuperar.")
add_body(doc, "Este livro existe para garantir que você não perca essa segunda janela que está aberta agora — enquanto a maioria ainda está em modo de observação.")
add_body(doc, "Não é um guia técnico. É um mapa de monetização.", italic=True)
add_body(doc, "Bom uso.")

page_break(doc)

# ── INTRODUÇÃO ────────────────────────────────────────────────────────────────

add_heading(doc, "Introdução — O Momento que Muda Carreiras", 1)

add_stat(doc, "US$ 1,8 trilhão", "projeção do mercado global de I.A. até 2030")
add_stat(doc, "< 5%", "dos profissionais criativos brasileiros usa I.A. de forma produtiva")
add_stat(doc, "2025", "ainda é o ano dos early adopters")

add_body(doc, "Existe um padrão que se repete em todas as grandes ondas tecnológicas: uma janela curta, no início, onde quem entra colhe vantagens desproporcionais. Depois essa janela fecha.")
add_body(doc, "Aconteceu com o SEO (quem fez sites em 2005 dominou o Google por uma década). Com o Instagram (quem criou conta em 2012 construiu audiências de milhões). Com o YouTube (os primeiros canais de 2007 ainda vivem de royalties).")
add_body(doc, "Agora é a vez da I.A. E a janela ainda está aberta.")

add_callout(doc, "A pergunta não é 'a I.A. vai me substituir?'. A pergunta certa é: 'como eu uso I.A. para entregar 5x mais valor e cobrar o dobro?'", "PONTO-CHAVE")

add_body(doc, "Este eBook não vai te ensinar a programar modelos de linguagem. Vai te mostrar quais ferramentas usar, para qual tipo de trabalho, e quanto cobrar por isso — com caminhos específicos para designers, cineastas, criadores de conteúdo, consultores e produtores de infoprodutos.")

page_break(doc)

# ── CAPÍTULO 1 ────────────────────────────────────────────────────────────────

add_heading(doc, "Capítulo 1 — O Mapa do Mercado de I.A. em 2025", 1)
add_heading(doc, "Três Camadas do Ecossistema", 2)

add_body(doc, "O mercado de I.A. pode ser dividido em três camadas que funcionam de forma complementar:")

add_bullet(doc, "Camada 1 — Modelos de Linguagem (LLMs): Claude, ChatGPT, Gemini. Processam texto, raciocinam, escrevem, planejam.")
add_bullet(doc, "Camada 2 — Geradores de Mídia: DALL-E 3, Imagen 3, Midjourney (imagem); Runway, Higgsfield.ai, Kling AI (vídeo); ElevenLabs, Suno (áudio).")
add_bullet(doc, "Camada 3 — Hubs e Amplificadores: Magnific.ai (upscale/refinamento), Higgsfield.ai (controle cinematográfico), Adobe Firefly (integrado ao ecossistema criativo).")

add_heading(doc, "Por Que Agora?", 2)
add_body(doc, "Em 2023, as ferramentas ainda eram instáveis e os resultados inconsistentes. Em 2024, chegaram à qualidade comercial. Em 2025, chegaram à acessibilidade de uso — qualquer profissional pode aprender o essencial em dias, não meses.")
add_body(doc, "O gap entre quem usa e quem não usa é agora visível nos orçamentos, na velocidade de entrega e na qualidade do portfólio.")

add_callout(doc, "Uma imagem gerada com Midjourney + refinada no Magnific.ai é indistinguível de uma foto de estúdio profissional. Um video criado no Higgsfield.ai com o prompt certo parece produção de R$ 50.000.", "EXEMPLO REAL")

add_heading(doc, "O Que o Mercado Está Comprando", 2)

services = [
    ("Design de marca com I.A.", "R$ 800 – R$ 3.000 por projeto"),
    ("Vídeos publicitários 15s/30s", "R$ 1.200 – R$ 5.000 por vídeo"),
    ("Conteúdo para redes sociais (pacote mensal)", "R$ 1.500 – R$ 4.000/mês"),
    ("Consultoria de implementação de I.A.", "R$ 3.000 – R$ 15.000 por empresa"),
    ("Cursos e infoprodutos sobre I.A.", "R$ 197 – R$ 2.000 por produto"),
]

for service, price in services:
    p = doc.add_paragraph()
    run_s = p.add_run(f"{service}: ")
    set_font(run_s, "Calibri", 11, bold=True)
    run_p = p.add_run(price)
    set_font(run_p, "Calibri", 11, color=(37, 99, 235))
    p.paragraph_format.space_after = Pt(4)

page_break(doc)

# ── CAPÍTULO 2 ────────────────────────────────────────────────────────────────

add_heading(doc, "Capítulo 2 — Os Hubs: Magnific.ai e Higgsfield.ai", 1)

add_heading(doc, "Magnific.ai — O Amplificador de Qualidade", 2)
add_body(doc, "Magnific.ai não gera imagens do zero. Ele pega uma imagem existente — seja gerada por I.A. ou fotografada — e a recria em resolução ultra-alta, adicionando detalhes fotorrealistas que o modelo original não capturou.")
add_body(doc, "Casos de uso profissional:")
add_bullet(doc, "Upscale de logos e artes para impressão em grande formato")
add_bullet(doc, "Refinamento de imagens geradas por Midjourney para uso comercial")
add_bullet(doc, "Restauração de fotos de produtos com baixa resolução")
add_bullet(doc, "Criação de mockups ultra-realistas de embalagens")

add_body(doc, "Quanto cobrar: R$ 150–800 por batch de imagens refinadas. Clientes: agências, e-commerces, marcas de moda.")

add_heading(doc, "Higgsfield.ai — O Estúdio Cinematográfico em Nuvem", 2)
add_body(doc, "Higgsfield.ai é o único gerador de vídeo com I.A. que oferece controle granular de câmera: você define o movimento (pan, tilt, dolly, zoom), a velocidade, o ângulo e o foco. O resultado é indistinguível de vídeo filmado com equipamento profissional.")
add_body(doc, "Casos de uso profissional:")
add_bullet(doc, "Vídeos publicitários para marcas (15s e 30s)")
add_bullet(doc, "Trailers e teasers de produtos digitais")
add_bullet(doc, "Conteúdo de marca para TikTok e Reels")
add_bullet(doc, "Vídeos de apresentação para pitch decks")

add_callout(doc, "Um vídeo publicitário de 30 segundos criado no Higgsfield.ai que antes custava R$ 8.000 em produção pode ser entregue por R$ 2.000 com qualidade superior. Margem do profissional: 70%+.", "OPORTUNIDADE")

page_break(doc)

# ── CAPÍTULO 3 ────────────────────────────────────────────────────────────────

add_heading(doc, "Capítulo 3 — Os Agentes: Claude, ChatGPT e Gemini", 1)
add_body(doc, "Os modelos de linguagem são o cérebro da operação com I.A. Cada um tem características distintas que os tornam mais adequados para tipos específicos de trabalho.")

add_heading(doc, "Claude (Anthropic)", 2)
add_body(doc, "Melhor para: escrita longa, análise de documentos, raciocínio complexo, código. O Claude se destaca em tarefas que exigem nuance e contexto extenso — ideal para criadores de conteúdo, consultores e produtores de infoprodutos.")
add_bullet(doc, "Limite de contexto: até 200k tokens (documentos inteiros)")
add_bullet(doc, "Ponto forte: consistência de voz e qualidade editorial")
add_bullet(doc, "Uso profissional: roteiros, relatórios, e-books, estratégias de marketing")

add_heading(doc, "ChatGPT (OpenAI)", 2)
add_body(doc, "Melhor para: versatilidade geral, geração de imagens (DALL-E 3 integrado), automações via GPTs customizados.")
add_bullet(doc, "DALL-E 3 integrado: geração de imagem diretamente na conversa")
add_bullet(doc, "GPTs: agentes personalizados para tarefas específicas")
add_bullet(doc, "Uso profissional: brainstorming visual, posts de redes sociais, atendimento automatizado")

add_heading(doc, "Gemini (Google)", 2)
add_body(doc, "Melhor para: integração com Google Workspace, análise de dados, pesquisa com acesso à web em tempo real.")
add_bullet(doc, "Integrado ao Google Docs, Gmail, Drive")
add_bullet(doc, "Imagen 3 integrado: qualidade de imagem de ponta")
add_bullet(doc, "Uso profissional: relatórios, análises de mercado, automação de e-mails")

add_heading(doc, "Como Escolher Qual Usar", 2)

choices = [
    ("Escrever e-book, roteiro, relatório", "Claude"),
    ("Gerar imagens + texto no mesmo fluxo", "ChatGPT"),
    ("Pesquisar dados atuais + planilhas", "Gemini"),
    ("Análise de documento longo (contrato, PDF)", "Claude"),
    ("Criar agente automatizado", "ChatGPT (GPTs)"),
]

for task, tool in choices:
    p = doc.add_paragraph()
    run_t = p.add_run(f"{task}  →  ")
    set_font(run_t, "Calibri", 11)
    run_tool = p.add_run(tool)
    set_font(run_tool, "Calibri", 11, bold=True, color=(37, 99, 235))
    p.paragraph_format.space_after = Pt(4)

page_break(doc)

# ── CAPÍTULO 4 ────────────────────────────────────────────────────────────────

add_heading(doc, "Capítulo 4 — Geradores de Imagem: DALL-E 3 e Imagen 3", 1)
add_body(doc, "A geração de imagem com I.A. passou de curiosidade para ferramenta comercial viável em 2024. Em 2025, a discussão não é mais 'I.A. pode fazer isso?' — é 'qual ferramenta gera o melhor resultado para este caso de uso específico?'")

add_heading(doc, "DALL-E 3 (OpenAI)", 2)
add_body(doc, "Integrado ao ChatGPT, o DALL-E 3 se destaca pela fidelidade a prompts textuais complexos — ele segue instruções detalhadas de composição, estilo e elementos com mais precisão que gerações anteriores.")
add_bullet(doc, "Ponto forte: seguimento de prompt, texto em imagem")
add_bullet(doc, "Uso ideal: ilustrações para e-books, posts com texto visual, thumbnails")
add_bullet(doc, "Preço: incluso no ChatGPT Plus ($20/mês)")

add_heading(doc, "Imagen 3 (Google)", 2)
add_body(doc, "O Imagen 3 lidera em fotorrealismo e qualidade de detalhes de pele e textura. Para fotografias de produto, retratos e cenas naturalistas, é a escolha superior.")
add_bullet(doc, "Ponto forte: fotorrealismo, texturas, iluminação natural")
add_bullet(doc, "Uso ideal: fotos de produto, retratos de marca, cenas ambientadas")
add_bullet(doc, "Preço: disponível via Google AI Studio e Gemini Advanced")

add_heading(doc, "Midjourney — O Padrão Criativo", 2)
add_body(doc, "Para arte conceitual, design editorial e imagens com forte identidade visual, o Midjourney ainda é o padrão de qualidade do mercado criativo. A curva de aprendizado do prompting vale o investimento.")
add_bullet(doc, "Ponto forte: qualidade artística, coerência de estilo")
add_bullet(doc, "Uso ideal: identidade visual, arte conceitual, editorial de moda")
add_bullet(doc, "Preço: a partir de $10/mês")

add_callout(doc, "Fluxo profissional de imagem: ChatGPT (DALL-E 3 ou Midjourney) para geração → Magnific.ai para refinamento e upscale → Entrega ao cliente. Tempo total: 20–40 minutos. Valor cobrado: R$ 300–1.500.", "FLUXO RECOMENDADO")

page_break(doc)

# ── CAPÍTULO 5 ────────────────────────────────────────────────────────────────

add_heading(doc, "Capítulo 5 — Carreiras e Quanto Cobrar", 1)
add_body(doc, "A pergunta mais comum de quem começa com I.A. é: 'mas o que exatamente eu vendo?' A resposta depende do seu perfil profissional. Abaixo, os caminhos mais lucrativos por tipo de profissional.")

careers = [
    (
        "Designer / Ilustrador",
        [
            "Identidade visual com I.A. (logo + paleta + tipografia): R$ 800–2.500",
            "Pack de redes sociais mensal (20 peças): R$ 1.200–3.000/mês",
            "Ilustrações para e-book ou apresentação: R$ 500–1.500",
            "Mockups de produto ultra-realistas: R$ 300–800 por batch",
        ],
        "Ferramenta principal: Midjourney + Magnific.ai + Canva Pro",
    ),
    (
        "Cineasta / Diretor de Vídeo",
        [
            "Vídeo publicitário 30s para marca: R$ 1.500–5.000",
            "Trailer de produto digital: R$ 800–2.500",
            "Pack de Reels/TikToks mensais (8 vídeos): R$ 2.000–4.500/mês",
            "Vídeo institucional 2 min: R$ 3.000–8.000",
        ],
        "Ferramenta principal: Higgsfield.ai + Runway + CapCut",
    ),
    (
        "Criador de Conteúdo",
        [
            "Gestão de perfil no Instagram/TikTok: R$ 1.500–4.000/mês",
            "Criação de roteiros + edição com I.A.: R$ 500–1.500 por vídeo",
            "Newsletter semanal para marca: R$ 800–2.000/mês",
            "Ghostwriting de threads e posts: R$ 300–800/post",
        ],
        "Ferramenta principal: Claude/ChatGPT + Kapwing + ElevenLabs",
    ),
    (
        "Consultor de Marketing / Estratégia",
        [
            "Diagnóstico de marketing com I.A. (relatório): R$ 2.500–8.000",
            "Implementação de fluxo de automação: R$ 5.000–15.000",
            "Treinamento de equipe em I.A.: R$ 3.000–10.000 por empresa",
            "Retainer mensal de consultoria: R$ 3.000–8.000/mês",
        ],
        "Ferramenta principal: Claude + Gemini + Make/Zapier",
    ),
    (
        "Produtor de Infoprodutos",
        [
            "E-book premium (design + conteúdo): R$ 47–197 por venda",
            "Curso gravado com I.A. (roteiro + slides + locução): R$ 497–2.000",
            "Mentoria em grupo: R$ 200–500/mês por aluno",
            "Certificação profissional: R$ 997–2.500",
        ],
        "Ferramenta principal: Claude + ChatGPT + ElevenLabs + Hotmart",
    ),
]

for career, services, tools in careers:
    add_heading(doc, career, 2)
    for s in services:
        add_bullet(doc, s)
    p = doc.add_paragraph()
    run = p.add_run(f"→ {tools}")
    set_font(run, "Calibri", 10, italic=True, color=(75, 85, 99))
    p.paragraph_format.space_after = Pt(10)

page_break(doc)

# ── CAPÍTULO 6 ────────────────────────────────────────────────────────────────

add_heading(doc, "Capítulo 6 — Fluxos de Trabalho Profissionais", 1)
add_body(doc, "A diferença entre amador e profissional em I.A. não é a ferramenta — é o fluxo. Profissionais conectam ferramentas em sequência, cada uma fazendo o que faz melhor, gerando um resultado que nenhuma faria sozinha.")

add_heading(doc, "Fluxo 1: Criação de Identidade Visual", 2)
steps1 = [
    ("Brief do cliente", "Conversa + Claude para organizar requisitos"),
    ("Geração de conceito visual", "Midjourney (10–20 variações de logo)"),
    ("Seleção e refinamento", "Magnific.ai (upscale) + Adobe Firefly (ajustes)"),
    ("Sistema de identidade", "Canva Pro (aplicação em templates)"),
    ("Entrega", "PDF + arquivos vetorizados"),
]
for step, tool in steps1:
    p = doc.add_paragraph()
    run_s = p.add_run(f"→ {step}: ")
    set_font(run_s, "Calibri", 11, bold=True)
    run_t = p.add_run(tool)
    set_font(run_t, "Calibri", 11)
    p.paragraph_format.space_after = Pt(3)

add_body(doc, "Tempo total: 4–8 horas (vs. 20–40h sem I.A.)  |  Margem: 75%+")

add_heading(doc, "Fluxo 2: Produção de Vídeo Publicitário", 2)
steps2 = [
    ("Roteiro e storyboard", "Claude (roteiro) + ChatGPT (variações de copy)"),
    ("Geração de cenas", "Higgsfield.ai (cenas principais) + Runway (transições)"),
    ("Locução e música", "ElevenLabs (voz) + Suno (trilha)"),
    ("Edição e montagem", "CapCut Pro ou Premiere"),
    ("Entrega", "MP4 em múltiplos formatos (16:9, 9:16, 1:1)"),
]
for step, tool in steps2:
    p = doc.add_paragraph()
    run_s = p.add_run(f"→ {step}: ")
    set_font(run_s, "Calibri", 11, bold=True)
    run_t = p.add_run(tool)
    set_font(run_t, "Calibri", 11)
    p.paragraph_format.space_after = Pt(3)

add_body(doc, "Tempo total: 6–12 horas (vs. 40–80h sem I.A.)  |  Valor: R$ 2.000–5.000")

add_heading(doc, "Fluxo 3: Produção de E-book Premium", 2)
steps3 = [
    ("Estrutura e outline", "Claude (estrutura de capítulos + argumentação)"),
    ("Rascunho de conteúdo", "Claude (escrita) + Gemini (pesquisa de dados)"),
    ("Design e diagramação", "ChatGPT+DALL-E 3 (imagens) + Canva/InDesign"),
    ("Revisão e edição", "Claude (revisão editorial)"),
    ("Entrega", "PDF + DOCX + HTML interativo"),
]
for step, tool in steps3:
    p = doc.add_paragraph()
    run_s = p.add_run(f"→ {step}: ")
    set_font(run_s, "Calibri", 11, bold=True)
    run_t = p.add_run(tool)
    set_font(run_t, "Calibri", 11)
    p.paragraph_format.space_after = Pt(3)

add_body(doc, "Tempo total: 8–16 horas (vs. semanas sem I.A.)  |  Valor: R$ 800–3.000 por e-book")

page_break(doc)

# ── CAPÍTULO 7 ────────────────────────────────────────────────────────────────

add_heading(doc, "Capítulo 7 — Seu Plano de 30 Dias", 1)
add_body(doc, "O maior erro de quem começa é tentar aprender tudo ao mesmo tempo. O plano abaixo é sequenciado para você faturar antes de dominar todas as ferramentas.")

add_heading(doc, "Semana 1: Fundamentos e Primeira Entrega", 2)
week1 = [
    "Escolha UMA ferramenta principal alinhada ao seu perfil (ver Cap. 5)",
    "Dedique 2h por dia aos prompts — use os exemplos do Bônus 1",
    "Crie 3 trabalhos de portfólio com I.A. (sem cliente, por sua conta)",
    "Monte um mini-portfólio no Canva ou Notion",
    "Ofereça 1 trabalho grátis para um contato para gerar depoimento",
]
for item in week1:
    add_bullet(doc, item)

add_heading(doc, "Semana 2: Primeiro Cliente e Primeiro Pagamento", 2)
week2 = [
    "Defina seu serviço-âncora (o mais simples, entregável em 48h)",
    "Precifique abaixo do mercado (50% do valor normal) para validar",
    "Ofereça para 5 potenciais clientes via DM, WhatsApp ou LinkedIn",
    "Entregue com agilidade e peça avaliação/depoimento",
    "Use o Bônus 2 para ajustar sua precificação com base no mercado",
]
for item in week2:
    add_bullet(doc, item)

add_heading(doc, "Semana 3–4: Sistematizar e Escalar", 2)
week3 = [
    "Crie templates de prompt para os serviços mais pedidos",
    "Desenvolva um pacote mensal recorrente (previsibilidade de caixa)",
    "Suba seu preço para o valor de mercado",
    "Adicione uma segunda ferramenta ao seu fluxo",
    "Comece a documentar resultados para marketing de conteúdo",
]
for item in week3:
    add_bullet(doc, item)

add_callout(doc, "Meta realista para o fim do mês 1: R$ 1.000–3.000 em serviços. Meta do mês 3, com portfólio e recorrência: R$ 3.000–8.000.", "META")

page_break(doc)

# ── CONCLUSÃO ─────────────────────────────────────────────────────────────────

add_heading(doc, "Conclusão — A Janela Está Aberta — Por Enquanto", 1)
add_body(doc, "Você chegou até aqui. Isso já te coloca à frente de 90% dos profissionais que vão passar os próximos 6 meses 'considerando entrar no mercado de I.A.'")
add_body(doc, "A janela existe agora. Não em 2026, não depois de fazer outro curso, não quando você 'se sentir pronto'. Agora.")
add_body(doc, "O mercado não espera. Os clientes que hoje pagam R$ 5.000 por um vídeo que você pode entregar por R$ 2.000 vão encontrar alguém que usa I.A. — a questão é se vai ser você.")

add_callout(doc, "\"A melhor hora para plantar uma árvore era há 20 anos. A segunda melhor hora é agora.\" O mesmo vale para I.A. em 2025.", "PARA REFLETIR")

add_body(doc, "Você tem o mapa. Você tem as ferramentas. Você tem os prompts.")
add_body(doc, "O próximo passo é seu.", italic=True)

page_break(doc)

# ── BÔNUS 1: 20 PROMPTS ──────────────────────────────────────────────────────

add_heading(doc, "Bônus 1 — 20 Prompts Prontos para Usar Hoje", 1)
add_body(doc, "Copie, adapte e use. Cada prompt foi testado para gerar resultados profissionais com o mínimo de iterações.")

prompts = [
    ("Design de Logo", "Crie um logo minimalista para [NOME DA MARCA], uma empresa de [SEGMENTO]. Estilo: moderno, clean, paleta [CORES]. Formato: ícone geométrico + logotipo. Variações: fundo branco e fundo escuro."),
    ("Post Instagram", "Escreva uma legenda para Instagram sobre [TEMA] para uma marca de [SEGMENTO]. Tom: [FORMAL/DESCONTRAÍDO]. Inclua: gancho na primeira linha, 3 blocos de valor, CTA, 10 hashtags relevantes. Máx 2.200 caracteres."),
    ("Vídeo Higgsfield.ai", "Cinematic shot of [DESCRIÇÃO DA CENA]. Camera: slow dolly forward, slight tilt up. Lighting: golden hour, warm tones. Style: commercial, premium brand feel. Duration: 5 seconds."),
    ("Identidade Visual", "Crie um sistema de identidade visual para [MARCA]. Inclua: paleta de 5 cores (hex), tipografia principal e secundária, regras de uso do logo, 3 exemplos de aplicação em materiais de marketing."),
    ("Roteiro de Vídeo 30s", "Escreva um roteiro de 30 segundos para anúncio de [PRODUTO/SERVIÇO]. Estrutura: 0-5s gancho (problema), 5-20s solução + prova, 20-30s CTA urgência. Tom: [EMPÁTICO/ENERGÉTICO]. Linguagem brasileira informal."),
    ("Copy de Anúncio Meta", "Escreva 3 variações de copy para anúncio Meta Ads de [PRODUTO]. Cada variação: headline (máx 40 char), texto principal (máx 125 char), descrição (máx 30 char). Ângulos diferentes: curiosidade, medo de perder, transformação."),
    ("Análise de Concorrente", "Analise a estratégia de marketing digital de [CONCORRENTE] com base nas informações a seguir: [COLAR DADOS]. Identifique: pontos fortes, fraquezas, oportunidades de diferenciação e gaps que posso explorar."),
    ("Proposta Comercial", "Escreva uma proposta comercial para [SERVIÇO] para o cliente [SEGMENTO DO CLIENTE]. Inclua: resumo executivo, escopo detalhado, cronograma, investimento (R$ [VALOR]), condições de pagamento, próximos passos."),
    ("Newsletter", "Escreva uma newsletter semanal sobre [TEMA] para uma lista de [PERFIL DA AUDIÊNCIA]. Inclua: assunto do e-mail (máx 50 char), preview text, 3 seções de conteúdo com títulos, CTA final. Tom editorial, sem jargão."),
    ("Bio de Perfil", "Escreva uma bio profissional para [PLATAFORMA: Instagram/LinkedIn/Site] para [NOME], especialista em [ÁREA]. Destaque: resultado que entrega, diferencial, credibilidade, CTA. Máx 150 caracteres (Instagram) ou 3 parágrafos (LinkedIn)."),
    ("Refinamento Magnific.ai", "Upscale this image to 4K resolution. Enhance: skin texture, fabric details, environmental lighting. Style: editorial photography, commercial grade. Preserve: original composition and color palette. Add: subtle film grain."),
    ("Script de Pitch", "Escreva um pitch de 60 segundos para [SERVIÇO/PRODUTO]. Estrutura: problema (10s), solução (20s), prova (15s), oferta (10s), CTA (5s). Para apresentação oral — frases curtas, linguagem conversacional."),
    ("Estratégia de Conteúdo", "Crie um plano de conteúdo de 30 dias para [MARCA/PROFISSIONAL] no [INSTAGRAM/TIKTOK/LINKEDIN]. Inclua: 12 temas de posts, 4 tipos de conteúdo (educativo, inspiracional, entretenimento, vendas), frequência, melhores horários."),
    ("Imagem de Produto", "Product photography of [PRODUTO] on [SUPERFÍCIE/CENÁRIO]. Lighting: [STUDIO SOFT/NATURAL WINDOW/NEON]. Style: [MINIMALISTA/LUXUOSO/MODERNO]. Background: [COR/MATERIAL]. Ultra-detailed, commercial quality, 4K."),
    ("E-mail de Vendas", "Escreva um e-mail de vendas para [PRODUTO/SERVIÇO] enviado para leads que [AÇÃO DO LEAD]. Assunto (urgência), abertura (empatia), corpo (problema → solução → prova), CTA único, PS com escassez. Máx 400 palavras."),
    ("Relatório Executivo", "Com base nos dados a seguir: [COLAR DADOS], crie um relatório executivo de 1 página. Inclua: sumário executivo, 3 insights principais, 1 gráfico sugerido, recomendações priorizadas, próximos passos. Linguagem clara, sem jargão técnico."),
    ("Persona de Comprador", "Crie uma persona detalhada para [PRODUTO/SERVIÇO]. Inclua: nome, idade, profissão, rotina diária, dores específicas, desejos, canais que consome, objeções de compra e linguagem que usa para descrever seus problemas."),
    ("Fluxo de Automação", "Desenhe um fluxo de automação de e-mail para [OBJETIVO: lead nurturing/pós-compra/reativação]. Inclua: trigger inicial, 5 e-mails com intervalo e assunto, segmentação por comportamento, condições de saída. Plataforma: [ACTIVECAMP/MAILCHIMP/RD]."),
    ("Título de Infoproduto", "Crie 10 variações de título para um [TIPO: e-book/curso/mentoria] sobre [TEMA]. Cada variação com: título principal (máx 8 palavras), subtítulo (máx 15 palavras). Ângulos: resultado, velocidade, especificidade, medo de perder."),
    ("Depoimento Solicitado", "Escreva um template de mensagem para pedir depoimento a clientes satisfeitos com [SERVIÇO]. Tom: natural, sem pressão. Inclua: agradecimento, contexto do pedido, 3 perguntas-guia, instruções simples de onde publicar."),
]

for i, (title, prompt) in enumerate(prompts, 1):
    add_heading(doc, f"Prompt {i:02d} — {title}", 3)
    add_body(doc, prompt, indent=True)

page_break(doc)

# ── BÔNUS 2: TABELA DE PREÇOS ─────────────────────────────────────────────────

add_heading(doc, "Bônus 2 — Tabela de Preços por Serviço", 1)
add_body(doc, "Referência de precificação baseada em valores praticados por profissionais que usam I.A. no mercado brasileiro em 2025.")

pricing = [
    ("DESIGN", [
        ("Logo + manual de marca simplificado", "R$ 800", "R$ 2.500"),
        ("Identidade visual completa", "R$ 2.500", "R$ 8.000"),
        ("Pack de posts mensais (20 peças)", "R$ 1.200", "R$ 3.500"),
        ("Mockups de produto (5 ângulos)", "R$ 400", "R$ 1.200"),
        ("Apresentação corporativa (20 slides)", "R$ 800", "R$ 2.000"),
    ]),
    ("VÍDEO", [
        ("Reels/TikTok (por vídeo)", "R$ 300", "R$ 800"),
        ("Pack mensal de 8 vídeos", "R$ 2.000", "R$ 5.000"),
        ("Anúncio publicitário 30s", "R$ 1.500", "R$ 5.000"),
        ("Vídeo institucional 2 min", "R$ 3.000", "R$ 8.000"),
        ("Trailer de produto/curso", "R$ 1.000", "R$ 3.000"),
    ]),
    ("CONTEÚDO", [
        ("Gestão de redes sociais (mensal)", "R$ 1.500", "R$ 4.500"),
        ("Roteiro de vídeo (por roteiro)", "R$ 300", "R$ 800"),
        ("Newsletter semanal (mensal)", "R$ 800", "R$ 2.000"),
        ("E-book (criação completa)", "R$ 1.500", "R$ 5.000"),
        ("Ghostwriting de threads/posts", "R$ 300", "R$ 800"),
    ]),
    ("CONSULTORIA", [
        ("Diagnóstico de marketing (relatório)", "R$ 2.500", "R$ 8.000"),
        ("Implementação de automação", "R$ 5.000", "R$ 15.000"),
        ("Treinamento de equipe (por empresa)", "R$ 3.000", "R$ 10.000"),
        ("Retainer mensal", "R$ 3.000", "R$ 8.000"),
        ("Estratégia digital (por projeto)", "R$ 4.000", "R$ 12.000"),
    ]),
]

for category, items in pricing:
    add_heading(doc, category, 2)
    for service, low, high in items:
        p = doc.add_paragraph()
        run_s = p.add_run(f"{service}: ")
        set_font(run_s, "Calibri", 11)
        run_r = p.add_run(f"{low} – {high}")
        set_font(run_r, "Calibri", 11, bold=True, color=(37, 99, 235))
        p.paragraph_format.space_after = Pt(3)

page_break(doc)

# ── BÔNUS 3: 30 FERRAMENTAS ──────────────────────────────────────────────────

add_heading(doc, "Bônus 3 — 30 Ferramentas Essenciais de I.A.", 1)

tools_list = [
    ("LINGUAGEM / TEXTO", [
        ("Claude (Anthropic)", "Escrita, análise, raciocínio complexo", "Gratuito / $20/mês"),
        ("ChatGPT (OpenAI)", "Versatilidade, GPTs, DALL-E 3", "Gratuito / $20/mês"),
        ("Gemini (Google)", "Integração Google, Imagen 3", "Gratuito / $20/mês"),
        ("Perplexity AI", "Pesquisa com fontes em tempo real", "Gratuito / $20/mês"),
        ("Notion AI", "Escrita integrada ao Notion", "$10/mês"),
    ]),
    ("IMAGEM", [
        ("Midjourney", "Arte conceitual, design editorial", "$10/mês"),
        ("DALL-E 3", "Seguimento de prompt, texto em imagem", "Incluso no ChatGPT Plus"),
        ("Imagen 3", "Fotorrealismo, texturas", "Via Gemini Advanced"),
        ("Adobe Firefly", "Integração Creative Cloud", "Incluso no Adobe CC"),
        ("Magnific.ai", "Upscale e refinamento de imagem", "$39/mês"),
        ("Ideogram", "Texto em imagens, tipografia", "Gratuito / $8/mês"),
        ("Leonardo.ai", "Consistência de personagem, gaming", "Gratuito / $12/mês"),
    ]),
    ("VÍDEO", [
        ("Higgsfield.ai", "Controle cinematográfico de câmera", "$19/mês"),
        ("Runway Gen-3", "Geração e edição de vídeo", "$15/mês"),
        ("Kling AI", "Vídeo de alta qualidade, movimentos naturais", "Créditos grátis"),
        ("Pika", "Vídeos curtos com estilo e efeitos", "Créditos grátis"),
        ("CapCut Pro", "Edição com I.A. integrada", "Gratuito / $8/mês"),
        ("Pika Labs", "Animação de imagens, vídeos curtos", "Gratuito / $8/mês"),
    ]),
    ("ÁUDIO", [
        ("ElevenLabs", "Síntese de voz ultra-realista", "Gratuito / $5/mês"),
        ("Suno", "Geração de música com letra", "Gratuito / $10/mês"),
        ("Udio", "Música instrumental e com vocal", "Gratuito / $10/mês"),
        ("Adobe Podcast", "Limpeza de áudio com I.A.", "Gratuito (beta)"),
    ]),
    ("AUTOMAÇÃO", [
        ("Make (Integromat)", "Automações visuais complexas", "Gratuito / $9/mês"),
        ("Zapier", "Integrações simples entre apps", "Gratuito / $20/mês"),
        ("n8n", "Automação open-source", "Gratuito (self-hosted)"),
        ("Clay", "Prospecção e enriquecimento de leads", "$149/mês"),
    ]),
    ("PRODUTIVIDADE", [
        ("Gamma", "Apresentações com I.A.", "Gratuito / $10/mês"),
        ("Tome", "Storytelling visual com I.A.", "Gratuito / $16/mês"),
        ("Descript", "Edição de vídeo/podcast por texto", "$12/mês"),
        ("Otter.ai", "Transcrição e resumo de reuniões", "Gratuito / $10/mês"),
    ]),
]

for category, tools in tools_list:
    add_heading(doc, category, 2)
    for name, desc, price in tools:
        p = doc.add_paragraph()
        run_n = p.add_run(f"{name}  ")
        set_font(run_n, "Calibri", 11, bold=True)
        run_d = p.add_run(f"— {desc}  ")
        set_font(run_d, "Calibri", 11)
        run_p = p.add_run(f"({price})")
        set_font(run_p, "Calibri", 10, color=(75, 85, 99))
        p.paragraph_format.space_after = Pt(3)

page_break(doc)

# ── AGRADECIMENTOS ────────────────────────────────────────────────────────────

add_heading(doc, "Agradecimentos", 1)
add_body(doc, "A todos os profissionais que compartilharam seus resultados, workflows e aprendizados nas comunidades de I.A. ao longo de 2023–2025. Vocês são a prova de que o mercado já existe — e está esperando por mais.")
add_body(doc, "E a você, leitor, por ter a coragem de agir enquanto a maioria ainda observa.")

add_divider(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ERA DA I.A. — Como Lucrar Agora Antes que Seja Tarde Demais")
set_font(run, "Calibri", 10, italic=True, color=(107, 114, 128))

# ── SALVAR ────────────────────────────────────────────────────────────────────

output_path = r"c:\Users\JHONATAN\Downloads\Marketing Digital\Produtos\mercado-ia-2025\mercado-ia-2025-ebook.docx"
doc.save(output_path)
print(f"OK DOCX gerado: {output_path}")
