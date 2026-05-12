"""
Gera DOCX para o order-bump "10 Projetos de I.A. Prontos".
"""

from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

section = doc.sections[0]
section.page_width = Mm(210)
section.page_height = Mm(297)
section.left_margin = Mm(25)
section.right_margin = Mm(25)
section.top_margin = Mm(22)
section.bottom_margin = Mm(22)

# paleta
OURO = (196, 146, 42)
POP = (245, 158, 11)
TEXTO = (30, 30, 30)
SUAVE = (85, 85, 85)


def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level=1, color=TEXTO):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    sizes = {1: 24, 2: 18, 3: 14, 4: 12}
    set_font(run, "Calibri", sizes.get(level, 12), bold=True, color=color)
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_body(doc, text, italic=False, bold=False, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, "Calibri", 11, italic=italic, bold=bold)
    if indent:
        p.paragraph_format.left_indent = Mm(8)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(16)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_font(run, "Calibri", 11)
    p.paragraph_format.left_indent = Mm(8 + level * 6)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_prompt_box(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Mm(8)
    p.paragraph_format.right_indent = Mm(4)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, "Consolas", 9, italic=True, color=SUAVE)
    return p


def add_callout(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Mm(8)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run_l = p.add_run(f"[{label}] ")
    set_font(run_l, "Calibri", 10, bold=True, color=OURO)
    run_t = p.add_run(text)
    set_font(run_t, "Calibri", 10, italic=True, color=SUAVE)
    return p


def add_stat_row(doc, time_val, price_val, tools_val):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    for label, val in [("Tempo", time_val), ("Valor", price_val), ("Ferramentas", tools_val)]:
        run_v = p.add_run(f"  {val} ")
        set_font(run_v, "Calibri", 16, bold=True, color=OURO)
        run_l = p.add_run(f"{label}  |")
        set_font(run_l, "Calibri", 8, color=SUAVE)
    return p


def add_divider(doc):
    p = doc.add_paragraph("─" * 72)
    run = p.runs[0]
    set_font(run, "Calibri", 9, color=OURO)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)


def add_resumo(doc, items):
    add_heading(doc, "Resumo", 3, color=OURO)
    for item in items:
        add_bullet(doc, item)


def page_break(doc):
    doc.add_page_break()


# === CAPA ===
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("10 PROJETOS DE I.A. PRONTOS")
set_font(run, "Calibri", 32, bold=True, color=TEXTO)
p.paragraph_format.space_before = Pt(60)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Prompts Revelados — Copie, Cole e Fature")
set_font(run, "Calibri", 18, italic=True, color=OURO)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("O kit que transforma qualquer iniciante em um\nprestador de serviços de I.A. em 24 horas")
set_font(run, "Calibri", 12, color=SUAVE)

add_divider(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Era da I.A. — 1ª Edição — 2025")
set_font(run, "Calibri", 10, color=SUAVE)

page_break(doc)

# === DIREITOS + AVISO ===
add_heading(doc, "Direitos Autorais", 2)
add_body(doc, "© 2025 — Todos os direitos reservados. Nenhuma parte deste material pode ser reproduzida sem autorização por escrito do autor.")
add_heading(doc, "Aviso Legal", 2)
add_body(doc, "Os resultados apresentados são exemplos reais, porém não constituem garantia de ganhos. Os resultados dependem de fatores individuais como dedicação, habilidade de execução e condições de mercado.")

page_break(doc)

# === SUMARIO ===
add_heading(doc, "Sumário", 1)
chapters = [
    ("Projeto 01", "Identidade Visual com I.A."),
    ("Projeto 02", "Copy de Vendas"),
    ("Projeto 03", "Roteiro para Reels"),
    ("Projeto 04", "Pack de Posts (30 dias)"),
    ("Projeto 05", "E-mail de Vendas"),
    ("Projeto 06", "Proposta Comercial"),
    ("Projeto 07", "Thumbnail YouTube"),
    ("Projeto 08", "Personal Brand"),
    ("Projeto 09", "Descrição E-commerce"),
    ("Projeto 10", "Vídeo Explicativo (Script)"),
    ("→", "Quanto Cobrar por Cada Projeto"),
    ("→", "Como Fechar os Primeiros Clientes"),
    ("→", "Próximos Passos"),
]
for prefix, title in chapters:
    p = doc.add_paragraph()
    run_p = p.add_run(f"{prefix}  ")
    set_font(run_p, "Calibri", 11, bold=True, color=OURO)
    run_t = p.add_run(title)
    set_font(run_t, "Calibri", 11)
    p.paragraph_format.space_after = Pt(4)

page_break(doc)

# === COMO USAR ===
add_heading(doc, "Como Usar Este Material", 1)
add_body(doc, "Imagine entrar em uma reunião com um cliente amanhã cedo, apresentar um resultado concreto ainda nesta tarde, e receber um Pix antes do anoitecer.")
add_body(doc, "Este material é diferente de tudo que você já viu sobre I.A. Aqui você recebe projetos completos — com o prompt exato que você vai copiar, colar e entregar.")
add_body(doc, "Não leia tudo de uma vez. Escolha 2 projetos, execute hoje, sinta a confiança de entregar algo real.", bold=True)
add_callout(doc, "Ferramentas", "ChatGPT (gratuito), Claude.ai (gratuito), Canva (gratuito), Midjourney ou Leonardo.AI. Custo para começar: R$0,00.")

page_break(doc)

# === PROJETOS ===

projetos = [
    {
        "num": "01", "titulo": "Identidade Visual com I.A.",
        "tempo": "2h", "valor": "R$497", "ferramentas": "3",
        "intro": "Toda empresa precisa de uma identidade visual. Designers cobram de R$800 a R$5.000. Com I.A., você entrega em 2 horas.",
        "passos": [
            ("Passo 1 — Briefing (15 min)", "Colete: nome da empresa, segmento, público-alvo, 3 adjetivos da marca, 3 marcas referência, cor que nunca usar."),
            ("Passo 2 — Prompt de Estratégia", "Você é um estrategista de marca. Crie identidade visual para: [NOME], segmento [NICHO], público [PÚBLICO], personalidade [3 ADJETIVOS]. Entregue: conceito central, paleta 5 cores HEX, tipografia Google Fonts, tom de voz."),
            ("Passo 3 — Logo no Midjourney", "Minimalist logo for [NOME], [SEGMENTO], [CONCEITO], primary color [HEX], clean geometric, white background, vector style, no text --ar 1:1 --style raw"),
            ("Passo 4 — Montagem Canva", "Busque 'Brand Guidelines PDF' no Canva. Monte: logo + paleta + tipografia + tom de voz + aplicações (cartão, post, anúncio)."),
        ],
        "resumo": ["Briefing: 15 min — Estratégia: 5 min", "Logo: Midjourney 30 min — PDF Canva: 1h", "Valor a cobrar: R$297 a R$497"],
    },
    {
        "num": "02", "titulo": "Copy de Vendas",
        "tempo": "1h", "valor": "R$397", "ferramentas": "1",
        "intro": "Copy é o texto que vende. Copywriters cobram R$500 a R$3.000 por uma página. Você entrega em 1 hora.",
        "passos": [
            ("Passo 1 — Briefing (10 min)", "Colete: produto, preço, problema que resolve, público, objeções comuns, resultado prometido, depoimentos."),
            ("Passo 2 — Prompt PAS + AIDA", "Copywriter especializado. Copy completa para: PRODUTO [NOME], PREÇO [VALOR], PROBLEMA [X], PÚBLICO [Y]. Headline + subheadline + gancho + agitação + solução + 8 benefícios + quebra objeções + prova social + oferta + CTA + P.S."),
            ("Passo 3 — Variações A/B", "Gere: 3 headlines alternativas, 2 CTAs alternativos, 1 versão curta (300 palavras) para anúncios pagos."),
            ("Passo 4 — Revisão", "Leia em voz alta. Elimine frases genéricas. Substitua adjetivos vazios por dados específicos."),
        ],
        "resumo": ["Briefing: 10 min — Copy: 10 min", "Variações A/B: 5 min — Revisão: 15 min", "Valor: R$197 a R$397 por entrega"],
    },
    {
        "num": "03", "titulo": "Roteiro para Reels",
        "tempo": "30min (4 roteiros)", "valor": "R$597/mês", "ferramentas": "1",
        "intro": "Reels com roteiro bem feito têm 3 a 10 vezes mais alcance. Você entrega 4 roteiros por semana.",
        "passos": [
            ("Passo 1 — Perfil do criador", "Crie roteiros para: CRIADOR [NOME], NICHO [X], PÚBLICO [Y], TOM [DESCONTRAÍDO/SÉRIO], OBJETIVO [VENDER/EDUCAR]."),
            ("Passo 2 — Pauta semanal", "4 temas: 1 Educativo + 1 Curiosidade + 1 Posicionamento + 1 Vendas."),
            ("Passo 3 — Roteiro completo", "GANCHO (0-3s) + CORPO (4-45s) + VIRADA (46-55s) + CTA (56-60s). Formato: [tempo] fala + indicação visual."),
        ],
        "resumo": ["Briefing: 15 min (só uma vez)", "Pauta: 5 min — 4 roteiros: 20 min", "Valor: R$297 a R$597/mês"],
    },
    {
        "num": "04", "titulo": "Pack de Posts (30 Dias)",
        "tempo": "2h", "valor": "R$397/mês", "ferramentas": "2",
        "intro": "Calendário completo com 30 posts prontos — legenda, hashtags e formato visual.",
        "passos": [
            ("Passo 1 — Estratégia", "Calendário 30 posts para [EMPRESA], [NICHO]. 40% educativo, 30% conexão, 20% prova social, 10% venda."),
            ("Passo 2 — Legendas em lote", "Para posts 1-10: legenda com gancho (120 chars), corpo, CTA, 10 hashtags. Repita para 11-20 e 21-30."),
            ("Passo 3 — Planilha compartilhável", "Google Sheets: Dia | Formato | Tema | Status | Link da arte."),
        ],
        "resumo": ["Calendário: 10 min — 30 legendas: 40 min", "Montagem: 15 min", "Valor: R$197 a R$397/mês"],
    },
    {
        "num": "05", "titulo": "E-mail de Vendas",
        "tempo": "1h (5 e-mails)", "valor": "R$597", "ferramentas": "1",
        "intro": "Uma sequência de e-mails bem escrita pode gerar mais vendas do que qualquer anúncio pago.",
        "passos": [
            ("Passo 1 — Mapa da sequência", "Sequência 5 e-mails para [PRODUTO]. Para cada: objetivo, assunto, tom, gatilho."),
            ("Passo 2 — Redação individual", "E-mail [N]. Estrutura: 2 assuntos + pré-header + saudação + hook + corpo + CTA. 300-500 palavras."),
            ("Passo 3 — Linhas de assunto", "10 opções: 2 curiosidade, 2 número, 2 urgência, 2 benefício, 2 personalização."),
        ],
        "resumo": ["Mapeamento: 10 min — 5 e-mails: 40 min", "10 linhas de assunto: 5 min", "Valor: R$297 a R$597 por sequência"],
    },
    {
        "num": "06", "titulo": "Proposta Comercial",
        "tempo": "1h", "valor": "R$397", "ferramentas": "2",
        "intro": "Uma proposta profissional pode triplicar a taxa de fechamento. Quem envia PDF já está à frente de 90%.",
        "passos": [
            ("Passo 1 — Briefing pós-reunião", "Contexto: PRESTADOR [nome, serviço, resultados]. CLIENTE [nome, segmento, problema, objetivo]."),
            ("Passo 2 — Redação", "10 seções: Capa + Sumário Executivo + Diagnóstico + Solução + Metodologia + Cases + Investimento + Próximos Passos + Sobre + Contato."),
            ("Passo 3 — Montagem Canva", "Template 'Proposta Comercial'. PDF 8-12 páginas + versão editável."),
        ],
        "resumo": ["Briefing: 10 min — Redação: 30 min", "Montagem: 15-20 min", "Valor: R$197 a R$397 por proposta"],
    },
    {
        "num": "07", "titulo": "Thumbnail YouTube",
        "tempo": "20min", "valor": "R$97 un.", "ferramentas": "2",
        "intro": "Thumbnail decide se alguém clica ou passa. Você resolve isso em 20 minutos.",
        "passos": [
            ("Passo 1 — Briefing", "Título do vídeo, emoção, cor do canal, foto do criador."),
            ("Passo 2 — Conceito visual", "Canvas 1280x720px. Foto + fundo I.A. + texto alto impacto (Bebas Neue). Teste: reduza para 120x67px."),
        ],
        "resumo": ["Briefing 5 min — Conceito 5 min — Canva 15 min", "Valor: R$47 a R$97 por thumbnail | R$497 pacote 10"],
    },
    {
        "num": "08", "titulo": "Personal Brand",
        "tempo": "1h", "valor": "R$797", "ferramentas": "2",
        "intro": "Kit completo de posicionamento: proposta de valor + tagline + pilares de conteúdo + bio em 3 formatos.",
        "passos": [
            ("Passo 1 — Mapeamento", "PROFISSIONAL: [NOME], [Aacute;REA], [EXPERIEcirc;NCIA], [RESULTADOS], [PÚBLICO], [DIFERENCIAL]."),
            ("Passo 2 — Posicionamento", "Proposta de valor (1 frase) + Tagline (8 palavras) + Mensagem central + 4 pilares + Antítese + História."),
            ("Passo 3 — Bios em 3 formatos", "Instagram (150 chars) + LinkedIn (300 palavras) + Curta para apresentações (50 palavras)."),
        ],
        "resumo": ["Mapeamento 10 min — Desenvolvimento 20 min — Bios 15 min", "Valor: R$397 a R$797 por kit completo"],
    },
    {
        "num": "09", "titulo": "Descrição E-commerce",
        "tempo": "20min/produto", "valor": "R$97 un.", "ferramentas": "1",
        "intro": "Descrição fraca é dinheiro deixado na mesa. Shopee, ML, Amazon perdem vendas por textos genéricos.",
        "passos": [
            ("Passo 1 — Briefing", "Produto, specs, público, problema, diferencial, plataforma."),
            ("Passo 2 — Descrição completa", "Título SEO + subtítulo + 4 parágrafos + 8 bullets + tabela specs + 5 P+R + 15 keywords."),
            ("Passo 3 — Versão anúncio", "Headline Meta (40 chars) + texto primário (125 chars) + descrição Google Shopping (150 chars)."),
        ],
        "resumo": ["Briefing 5 min — Descrição 10 min — Anúncios 5 min", "Valor: R$47 a R$97 por produto | R$797 lote 10"],
    },
    {
        "num": "10", "titulo": "Vídeo Explicativo (Script)",
        "tempo": "1h", "valor": "R$597", "ferramentas": "2",
        "intro": "Explainer videos são usados em landing pages e redes sociais. O roteiro é a parte mais difícil.",
        "passos": [
            ("Passo 1 — Briefing", "EMPRESA [nome], objetivo [X], duração [60-120s], público [Y], problema [Z], solução [W]."),
            ("Passo 2 — Roteiro em tabela", "2 colunas: NARRAÇÃO + VISUAL. Ganchos 0-10s → Agitação → Solução → Resultado → CTA."),
            ("Passo 3 — Revisão de ritmo", "Marque 3 frases em destaque, identifique pausas dramáticas, sugira momentos de música."),
        ],
        "resumo": ["Briefing 10 min — Roteiro 25 min — Revisão 10 min", "Valor: R$297 a R$597 por roteiro"],
    },
]

for proj in projetos:
    add_heading(doc, f"Projeto {proj['num']} — {proj['titulo']}", 1)
    add_body(doc, proj["intro"])
    add_stat_row(doc, proj["tempo"], proj["valor"], proj["ferramentas"])
    add_divider(doc)

    for titulo, conteudo in proj["passos"]:
        add_heading(doc, titulo, 3)
        add_prompt_box(doc, conteudo)

    add_resumo(doc, proj["resumo"])
    page_break(doc)

# === QUANTO COBRAR ===
add_heading(doc, "Quanto Cobrar por Cada Projeto", 1)
add_body(doc, "Referência de precificação por nível de experiência:", italic=True)

tabela = [
    ("Identidade Visual", "R$197", "R$397", "R$697"),
    ("Copy de Vendas", "R$147", "R$297", "R$497"),
    ("Roteiros Reels (mês)", "R$197", "R$397", "R$597"),
    ("Pack 30 Posts (mês)", "R$147", "R$297", "R$497"),
    ("Sequência E-mails", "R$197", "R$397", "R$697"),
    ("Proposta Comercial", "R$147", "R$297", "R$497"),
    ("Thumbnail (unitário)", "R$47", "R$77", "R$127"),
    ("Personal Brand Kit", "R$297", "R$597", "R$997"),
    ("Descrição E-commerce", "R$37", "R$67", "R$97"),
    ("Roteiro Vídeo", "R$197", "R$397", "R$697"),
]

for nome, ini, inter, avanc in tabela:
    p = doc.add_paragraph()
    run_n = p.add_run(f"{nome}: ")
    set_font(run_n, "Calibri", 10)
    run_v = p.add_run(f"{ini} → {inter} → {avanc}")
    set_font(run_v, "Calibri", 10, bold=True, color=OURO)
    p.paragraph_format.space_after = Pt(3)

add_divider(doc)

add_heading(doc, "Pacotes Mensais", 2)
for nome, valor, desc in [
    ("Starter", "R$397/mês", "30 posts + 4 roteiros reels"),
    ("Crescimento", "R$697/mês", "30 posts + 4 reels + 1 copy"),
    ("Premium", "R$1.197/mês", "Tudo + thumbnails + e-mail marketing"),
]:
    p = doc.add_paragraph()
    run_n = p.add_run(f"{nome} ")
    set_font(run_n, "Calibri", 11, bold=True)
    run_v = p.add_run(f"{valor} ")
    set_font(run_v, "Calibri", 11, bold=True, color=OURO)
    run_d = p.add_run(f"— {desc}")
    set_font(run_d, "Calibri", 10, color=SUAVE)
    p.paragraph_format.space_after = Pt(4)

add_callout(doc, "Regra de ouro", "Nunca precifique pelo tempo. Precifique pelo resultado. Uma copy que gera R$10.000 em vendas vale R$500, não R$50.")

page_break(doc)

# === COMO FECHAR CLIENTES ===
add_heading(doc, "Como Fechar os Primeiros Clientes", 1)
add_heading(doc, "3 Canais sem Investimento", 2)

add_heading(doc, "1. Rede Pessoal (mais rápido)", 3)
add_body(doc, "Envie a mensagem abaixo para 10 pessoas que têm negócios:")
add_prompt_box(doc, "Oi [NOME], tudo bem? Estou expandindo meu trabalho com I.A. para ajudar negócios locais. Estou selecionando 3 clientes para fazer um projeto de [SERVIÇO] por um valor reduzido — em troca de feedback e depoimento. Você teria interesse? Leva menos de 48h para entregar.")

add_heading(doc, "2. Grupos WhatsApp e Facebook", 3)
add_body(doc, "Procure grupos de empreendedores. Não faça propaganda — responda perguntas e ofereça ajuda quando alguém tem o problema que você resolve.")

add_heading(doc, "3. LinkedIn (B2B)", 3)
add_body(doc, "Publique 1x/semana mostrando o processo de um projeto com antes/depois. Posts de processo têm alto engajamento.")

add_callout(doc, "Importante", "Nunca diga 'faço com I.A.'. Diga 'utilizo ferramentas de tecnologia avançada para entregar mais rápido com qualidade superior.'")

add_heading(doc, "Perguntas Frequentes", 2)
faqs = [
    ("Preciso saber programar?", "Não. Todos os projetos usam ferramentas no-code: ChatGPT, Canva, Midjourney."),
    ("E se o cliente pedir revisão?", "Inclua 2 rodadas de revisão. Ajustes em prompts são rápidos. Cobre R$50-100 por rodada extra."),
    ("Funciona para qualquer nicho?", "Sim. Os prompts são templates universais. Basta substituir as variáveis entre colchetes."),
]
for pergunta, resposta in faqs:
    add_heading(doc, pergunta, 4)
    add_body(doc, resposta)

page_break(doc)

# === PROXIMOS PASSOS ===
add_heading(doc, "Próximos Passos", 1)
add_body(doc, "Você tem os prompts. Tem as ferramentas. Tem os preços. O único ingrediente que falta é a execução.")
add_heading(doc, "Esta semana", 2)
steps = [
    "Escolha 1 projeto que mais faz sentido para o seu perfil",
    "Execute para você mesmo — crie sua própria identidade, roteiro ou proposta",
    "Documente o processo (print, vídeo curto, texto)",
    "Aborde 3 potenciais clientes com o projeto como case",
]
for i, step in enumerate(steps, 1):
    add_body(doc, f"{i}. {step}")

add_callout(doc, "Lembrete", "Você não precisa de portfólio antes de começar. O portfólio se constrói enquanto você trabalha.")

add_body(doc, "O mercado de I.A. está em formação agora. Os profissionais que se posicionarem nos próximos 12 meses vão colher resultados que os atrasados vão se arrepender de ter perdido.", italic=True)
add_body(doc, "A janela está aberta. Entre antes que feche.", bold=True)

page_break(doc)

# === SOBRE + AGRADECIMENTOS ===
add_heading(doc, "Sobre o Autor", 1)
add_body(doc, "Especialista em I.A. aplicada ao marketing digital, com experiência em criar sistemas de produção de conteúdo e automação para empreendedores e agências.")
add_body(doc, "Acredita que a I.A. não vai substituir profissionais — vai substituir os que não souberem usá-la.")

add_divider(doc)

add_heading(doc, "Agradecimentos", 1)
add_body(doc, "Obrigado por investir neste material. A decisão de comprar mostra que você é o tipo de pessoa que age — não apenas que observa.")
add_body(doc, "Agora é só executar.", bold=True)

add_divider(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("10 Projetos de I.A. Prontos — Era da I.A. — 2025")
set_font(run, "Calibri", 9, italic=True, color=SUAVE)

# === SALVAR ===
output = r"C:\Users\JHONATAN\Downloads\Skill Marketing Digital\Produtos\mercado-ia-2025\funil\order-bump\10-projetos-ia-prontos\10-projetos-ia-prontos.docx"
doc.save(output)
print(f"OK DOCX gerado: {output}")
